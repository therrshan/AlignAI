"""
Document processing module for extracting and structuring text from PDFs and DOCX files
Handles resume parsing, project extraction, and text chunking for vector storage
"""

import re
import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import PyPDF2
import pdfplumber
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from config.settings import RESUME_SECTIONS, CHUNK_SIZE, CHUNK_OVERLAP

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document extraction and processing for resumes and projects"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=CHUNK_SIZE,
            chunk_overlap=CHUNK_OVERLAP,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    def extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF using multiple methods for reliability"""
        text = ""
        
        # Method 1: Try pdfplumber first (better for formatted text)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if text.strip():
                logger.info(f"Successfully extracted text using pdfplumber: {len(text)} chars")
                return text.strip()
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}")
        
        # Method 2: Fallback to PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            if text.strip():
                logger.info(f"Successfully extracted text using PyPDF2: {len(text)} chars")
                return text.strip()
        except Exception as e:
            logger.error(f"PyPDF2 also failed: {e}")
        
        raise ValueError(f"Could not extract text from PDF: {file_path}")
    
    def extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.info(f"Successfully extracted text from DOCX: {len(text)} chars")
            return text.strip()
        
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {e}")
            raise ValueError(f"Could not extract text from DOCX: {file_path}")
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text from supported file formats"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif extension == '.docx':
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    def parse_resume_sections(self, resume_text: str) -> Dict[str, str]:
        """Parse resume into structured sections"""
        sections = {}
        
        # Convert text to lowercase for matching
        text_lower = resume_text.lower()
        
        # Find section boundaries
        section_positions = {}
        
        for section_key, section_headers in RESUME_SECTIONS.items():
            for header in section_headers:
                # Look for section headers (case insensitive)
                pattern = rf'\b{re.escape(header)}\b'
                match = re.search(pattern, text_lower)
                if match:
                    section_positions[section_key] = match.start()
                    break
        
        # Sort sections by position in document
        sorted_sections = sorted(section_positions.items(), key=lambda x: x[1])
        
        # Extract content between sections
        for i, (section_name, start_pos) in enumerate(sorted_sections):
            # Determine end position
            if i + 1 < len(sorted_sections):
                end_pos = sorted_sections[i + 1][1]
            else:
                end_pos = len(resume_text)
            
            # Extract section content
            section_content = resume_text[start_pos:end_pos].strip()
            
            # Clean up the content (remove the header line)
            lines = section_content.split('\n')
            if lines:
                # Remove the header line
                cleaned_lines = lines[1:] if len(lines) > 1 else lines
                sections[section_name] = '\n'.join(cleaned_lines).strip()
        
        # If no sections found, put everything in summary
        if not sections:
            sections['summary'] = resume_text
        
        logger.info(f"Parsed resume into {len(sections)} sections: {list(sections.keys())}")
        return sections
    
    def extract_projects_from_portfolio(self, projects_text: str) -> List[Dict[str, str]]:
        """Extract individual projects from a projects portfolio PDF"""
        projects = []
        
        # Split by common project delimiters
        project_patterns = [
            r'\n(?=\d+\.\s)',  # Numbered projects
            r'\n(?=[A-Z][^.]*(?:Project|System|Tool|Platform|Application))',  # Project titles
            r'\n(?=\w+\s*[-–—]\s*)',  # Project name - description format
        ]
        
        # Try different splitting patterns
        for pattern in project_patterns:
            project_sections = re.split(pattern, projects_text)
            if len(project_sections) > 1:
                break
        else:
            # If no pattern matches, split by double newlines
            project_sections = projects_text.split('\n\n')
        
        for i, section in enumerate(project_sections):
            section = section.strip()
            if len(section) > 100:  # Filter out very short sections
                # Extract project name (first line or before first dash/colon)
                lines = section.split('\n')
                first_line = lines[0].strip()
                
                # Clean up project name
                project_name = re.split(r'[-–—:]', first_line)[0].strip()
                project_name = re.sub(r'^\d+\.\s*', '', project_name)  # Remove numbering
                
                if not project_name:
                    project_name = f"Project {i+1}"
                
                projects.append({
                    'name': project_name,
                    'description': section,
                    'index': i
                })
        
        logger.info(f"Extracted {len(projects)} projects from portfolio")
        return projects
    
    def chunk_text_for_embedding(self, text: str, metadata: Dict = None) -> List[Dict]:
        """Split text into chunks suitable for embedding"""
        chunks = self.text_splitter.split_text(text)
        
        chunked_docs = []
        for i, chunk in enumerate(chunks):
            chunk_metadata = {
                'chunk_index': i,
                'chunk_size': len(chunk),
                **(metadata or {})
            }
            chunked_docs.append({
                'content': chunk,
                'metadata': chunk_metadata
            })
        
        logger.info(f"Created {len(chunked_docs)} chunks from text")
        return chunked_docs
    
    def process_resume(self, file_path: Path, resume_id: str) -> Tuple[str, Dict[str, str], List[Dict]]:
        """Complete resume processing pipeline"""
        logger.info(f"Processing resume: {file_path}")
        
        # Extract text
        full_text = self.extract_text(file_path)
        
        # Parse sections
        sections = self.parse_resume_sections(full_text)
        
        # Create chunks for embedding
        chunks = self.chunk_text_for_embedding(
            full_text, 
            metadata={
                'document_type': 'resume',
                'resume_id': resume_id,
                'file_name': file_path.name
            }
        )
        
        return full_text, sections, chunks
    
    def process_projects_portfolio(self, file_path: Path) -> Tuple[List[Dict], List[Dict]]:
        """Complete projects portfolio processing pipeline"""
        logger.info(f"Processing projects portfolio: {file_path}")
        
        # Extract text
        full_text = self.extract_text(file_path)
        
        # Extract individual projects
        projects = self.extract_projects_from_portfolio(full_text)
        
        # Create chunks for each project
        all_chunks = []
        for project in projects:
            project_chunks = self.chunk_text_for_embedding(
                project['description'],
                metadata={
                    'document_type': 'project',
                    'project_name': project['name'],
                    'project_index': project['index'],
                    'file_name': file_path.name
                }
            )
            all_chunks.extend(project_chunks)
        
        return projects, all_chunks

# Usage example and testing
if __name__ == "__main__":
    processor = DocumentProcessor()
    
    # Test with sample files (if they exist)
    test_files = ["/mnt/data/gitrepos/resume-ai/Darshan-Rahul-Rajopadhye.pdf"]
    
    for test_file in test_files:
        if Path(test_file).exists():
            try:
                text = processor.extract_text(Path(test_file))
                print(f"Successfully processed {test_file}: {len(text)} characters")
            except Exception as e:
                print(f"Error processing {test_file}: {e}")